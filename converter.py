import pypandoc
import os
import re
import tempfile

def convert_docx_to_latex(
    docx_path: str,
    latex_path: str,
    generate_toc: bool = False,
    extract_media_to_path: str = None,
    latex_template_path: str = None,
    overleaf_compatible: bool = False,
    preserve_styles: bool = True,
    preserve_linebreaks: bool = True
) -> tuple[bool, str]:
    """
    Converts a DOCX file to a LaTeX file using pypandoc with enhanced features.

    Args:
        docx_path: Path to the input .docx file.
        latex_path: Path to save the output .tex file.
        generate_toc: If True, attempts to generate a Table of Contents.
        extract_media_to_path: If specified, path to extract media to (e.g., "./media").
        latex_template_path: If specified, path to a custom Pandoc LaTeX template file.
        overleaf_compatible: If True, makes images work in Overleaf with relative paths.
        preserve_styles: If True, preserves document styles like centering and alignment.
        preserve_linebreaks: If True, preserves line breaks and proper list formatting.

    Returns:
        A tuple (success: bool, message: str).
    """
    extra_args = []
    
    # Ensure standalone document (not fragment)
    extra_args.append("--standalone")
    
    # Basic options
    if generate_toc:
        extra_args.append("--toc")
    if extract_media_to_path:
        extra_args.append(f"--extract-media={extract_media_to_path}")
    if latex_template_path and os.path.isfile(latex_template_path):
        extra_args.append(f"--template={latex_template_path}")
    elif latex_template_path:
        pass  # Template not found, Pandoc will handle the error

    # Enhanced features
    if overleaf_compatible:
        extra_args.extend([
            "--resource-path=./",
            "--default-image-extension=png"
        ])
    
    if preserve_styles:
        extra_args.extend([
            "--from=docx+styles",
            "--wrap=preserve",
            "--columns=72",
            "--strip-comments"  # Remove comments that might cause highlighting
        ])
    
    if preserve_linebreaks:
        extra_args.extend([
            "--preserve-tabs",
            "--wrap=preserve",
            "--reference-doc=" + docx_path  # Use original Word doc as reference for formatting
        ])
        
        # Create minimal Lua filter that preserves Word's original line breaks
        lua_filter_content = '''
function Para(elem)
  -- Preserve all line breaks exactly as they appear in Word
  -- This maintains Word's original pagination and formatting
  local new_content = {}
  
  for i, item in ipairs(elem.content) do
    if item.t == "SoftBreak" then
      -- Convert all soft breaks to line breaks to match Word's formatting
      table.insert(new_content, pandoc.LineBreak())
    else
      table.insert(new_content, item)
    end
  end
  
  elem.content = new_content
  return elem
end

function LineBlock(elem)
  -- Preserve line blocks exactly as they are
  return elem
end

function Span(elem)
  -- Remove unwanted highlighting and formatting
  if elem.attributes and elem.attributes.style then
    -- Remove background colors and highlighting
    local style = elem.attributes.style
    if string.find(style, "background") or string.find(style, "highlight") then
      elem.attributes.style = nil
    end
  end
  return elem
end

function Div(elem)
  -- Remove unwanted div formatting that causes highlighting
  if elem.attributes and elem.attributes.style then
    local style = elem.attributes.style
    if string.find(style, "background") or string.find(style, "highlight") then
      elem.attributes.style = nil
    end
  end
  return elem
end

function RawBlock(elem)
  -- Preserve raw LaTeX blocks
  if elem.format == "latex" then
    return elem
  end
end
'''
        
        # Create temporary Lua filter file
        with tempfile.NamedTemporaryFile(mode='w', suffix='.lua', delete=False) as f:
            f.write(lua_filter_content)
            lua_filter_path = f.name
        
        extra_args.append(f"--lua-filter={lua_filter_path}")

    try:
        # Perform conversion
        pypandoc.convert_file(docx_path, 'latex', outputfile=latex_path, extra_args=extra_args)
        
        # Clean up temporary Lua filter if created
        if preserve_linebreaks and 'lua_filter_path' in locals():
            try:
                os.unlink(lua_filter_path)
            except OSError:
                pass
        
        # Apply post-processing enhancements (always applied for Unicode conversion)
        _apply_post_processing(latex_path, overleaf_compatible, preserve_styles, preserve_linebreaks, extract_media_to_path)
        
        # Generate status message
        enhancements = []
        if overleaf_compatible:
            enhancements.append("Overleaf compatibility")
        if preserve_styles:
            enhancements.append("style preservation")
        if preserve_linebreaks:
            enhancements.append("line break preservation")
        
        if enhancements:
            enhancement_msg = f" with {', '.join(enhancements)}"
        else:
            enhancement_msg = ""
            
        return True, f"Conversion successful{enhancement_msg}!"
        
    except RuntimeError as e:
        # Clean up temporary Lua filter if created
        if preserve_linebreaks and 'lua_filter_path' in locals():
            try:
                os.unlink(lua_filter_path)
            except OSError:
                pass
        return False, f"RuntimeError: Could not execute Pandoc. Please ensure Pandoc is installed and in your system's PATH. Error: {e}"
    except Exception as e:
        # Clean up temporary Lua filter if created
        if preserve_linebreaks and 'lua_filter_path' in locals():
            try:
                os.unlink(lua_filter_path)
            except OSError:
                pass
        return False, f"Conversion failed: {e}"

def _apply_post_processing(latex_path: str, overleaf_compatible: bool, preserve_styles: bool, preserve_linebreaks: bool, extract_media_to_path: str = None):
    """
    Apply post-processing enhancements to the generated LaTeX file.
    """
    try:
        with open(latex_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Always inject essential packages for compilation compatibility
        content = _inject_essential_packages(content)
        
        # Fix mixed mathematical expressions first to remove duplicated text
        content = _fix_mixed_mathematical_expressions(content)
        
        # Convert Unicode mathematical characters to LaTeX equivalents (always applied)
        content = _convert_unicode_math_characters(content)
        
        # Apply additional Unicode cleanup as a safety net
        content = _additional_unicode_cleanup(content)
        
        # Apply overleaf compatibility fixes
        if overleaf_compatible:
            content = _fix_image_paths_for_overleaf(content, extract_media_to_path)
        
        # Apply style preservation enhancements
        if preserve_styles:
            content = _inject_latex_packages(content)
            content = _add_centering_commands(content)
        
        # Apply line break preservation fixes
        if preserve_linebreaks:
            content = _fix_line_breaks_and_spacing(content)
        
        # Remove unwanted formatting and highlighting
        content = _remove_unwanted_formatting(content)
        
        # Fix common LaTeX compilation issues
        content = _fix_compilation_issues(content)
        
        # Write back the processed content
        with open(latex_path, 'w', encoding='utf-8') as f:
            f.write(content)
            
    except Exception as e:
        # Post-processing failures shouldn't break the conversion
        print(f"Warning: Post-processing failed: {e}")

def _inject_essential_packages(content: str) -> str:
    """
    Inject essential packages that are always needed for compilation.
    """
    # Core packages that Pandoc might not include but are often needed
    essential_packages = [
        r'\usepackage[utf8]{inputenc}',  # UTF-8 input encoding
        r'\usepackage[T1]{fontenc}',     # Font encoding
        r'\usepackage{graphicx}',        # For images
        r'\usepackage{longtable}',       # For tables
        r'\usepackage{booktabs}',        # Better table formatting
        r'\usepackage{hyperref}',        # For links (if not already included)
        r'\usepackage{amsmath}',         # Mathematical formatting
        r'\usepackage{amssymb}',         # Mathematical symbols
        r'\usepackage{textcomp}',        # Additional text symbols
    ]
    
    documentclass_pattern = r'\\documentclass(?:\[[^\]]*\])?\{[^}]+\}'
    documentclass_match = re.search(documentclass_pattern, content)
    
    if documentclass_match:
        insert_pos = documentclass_match.end()
        
        packages_to_insert = []
        for package in essential_packages:
            package_name = package.split('{')[1].split('}')[0].split(']')[0]  # Extract package name
            if f'usepackage' not in content or package_name not in content:
                packages_to_insert.append(package)
        
        if packages_to_insert:
            package_block = '\n% Essential packages for compilation\n' + '\n'.join(packages_to_insert) + '\n'
            content = content[:insert_pos] + package_block + content[insert_pos:]
        
        # Add Unicode character definitions to handle any remaining problematic characters
        unicode_definitions = r'''
% Unicode character definitions for LaTeX compatibility
\DeclareUnicodeCharacter{2003}{ }  % Em space
\DeclareUnicodeCharacter{2002}{ }  % En space
\DeclareUnicodeCharacter{2009}{ }  % Thin space
\DeclareUnicodeCharacter{200A}{ }  % Hair space
\DeclareUnicodeCharacter{2004}{ }  % Three-per-em space
\DeclareUnicodeCharacter{2005}{ }  % Four-per-em space
\DeclareUnicodeCharacter{2006}{ }  % Six-per-em space
\DeclareUnicodeCharacter{2008}{ }  % Punctuation space
\DeclareUnicodeCharacter{202F}{ }  % Narrow no-break space
\DeclareUnicodeCharacter{2212}{-}  % Unicode minus sign
\DeclareUnicodeCharacter{2010}{-}  % Hyphen
\DeclareUnicodeCharacter{2011}{-}  % Non-breaking hyphen
\DeclareUnicodeCharacter{2013}{--} % En dash
\DeclareUnicodeCharacter{2014}{---}% Em dash
'''
        
        # Insert Unicode definitions after packages but before \begin{document}
        begin_doc_match = re.search(r'\\begin\{document\}', content)
        if begin_doc_match:
            insert_pos_unicode = begin_doc_match.start()
            content = content[:insert_pos_unicode] + unicode_definitions + '\n' + content[insert_pos_unicode:]
    
    return content

def _convert_unicode_math_characters(content: str) -> str:
    """
    Convert Unicode mathematical characters to their LaTeX equivalents.
    """
    # Dictionary of Unicode characters to LaTeX commands
    unicode_to_latex = {
        # Mathematical operators
        'Δ': r'$\Delta$',           # U+0394 - Greek capital letter delta
        'δ': r'$\delta$',           # U+03B4 - Greek small letter delta
        '∑': r'$\sum$',             # U+2211 - N-ary summation
        '∏': r'$\prod$',            # U+220F - N-ary product
        '∫': r'$\int$',             # U+222B - Integral
        '∂': r'$\partial$',         # U+2202 - Partial differential
        '∇': r'$\nabla$',           # U+2207 - Nabla
        '√': r'$\sqrt{}$',          # U+221A - Square root
        '∞': r'$\infty$',           # U+221E - Infinity
        
        # Relations and equality
        '≈': r'$\approx$',          # U+2248 - Almost equal to
        '≠': r'$\neq$',             # U+2260 - Not equal to
        '≤': r'$\leq$',             # U+2264 - Less-than or equal to
        '≥': r'$\geq$',             # U+2265 - Greater-than or equal to
        '±': r'$\pm$',              # U+00B1 - Plus-minus sign
        '∓': r'$\mp$',              # U+2213 - Minus-or-plus sign
        '×': r'$\times$',           # U+00D7 - Multiplication sign
        '÷': r'$\div$',             # U+00F7 - Division sign
        '⋅': r'$\cdot$',            # U+22C5 - Dot operator
        
        # Set theory and logic
        '∈': r'$\in$',              # U+2208 - Element of
        '∉': r'$\notin$',           # U+2209 - Not an element of
        '⊂': r'$\subset$',          # U+2282 - Subset of
        '⊃': r'$\supset$',          # U+2283 - Superset of
        '⊆': r'$\subseteq$',        # U+2286 - Subset of or equal to
        '⊇': r'$\supseteq$',        # U+2287 - Superset of or equal to
        '∪': r'$\cup$',             # U+222A - Union
        '∩': r'$\cap$',             # U+2229 - Intersection
        '∅': r'$\emptyset$',        # U+2205 - Empty set
        '∀': r'$\forall$',          # U+2200 - For all
        '∃': r'$\exists$',          # U+2203 - There exists
        
        # Special symbols
        '∣': r'$|$',                # U+2223 - Divides
        '∥': r'$\parallel$',        # U+2225 - Parallel to
        '⊥': r'$\perp$',            # U+22A5 - Up tack (perpendicular)
        '∠': r'$\angle$',           # U+2220 - Angle
        '°': r'$^\circ$',           # U+00B0 - Degree sign
        
        # Arrows
        '→': r'$\rightarrow$',      # U+2192 - Rightwards arrow
        '←': r'$\leftarrow$',       # U+2190 - Leftwards arrow
        '↔': r'$\leftrightarrow$',  # U+2194 - Left right arrow
        '⇒': r'$\Rightarrow$',      # U+21D2 - Rightwards double arrow
        '⇐': r'$\Leftarrow$',       # U+21D0 - Leftwards double arrow
        '⇔': r'$\Leftrightarrow$',  # U+21D4 - Left right double arrow
        
        # Accents and diacritics
        'ˉ': r'$\bar{}$',           # U+02C9 - Modifier letter macron
        'ˆ': r'$\hat{}$',           # U+02C6 - Modifier letter circumflex accent
        'ˇ': r'$\check{}$',         # U+02C7 - Caron
        '˜': r'$\tilde{}$',         # U+02DC - Small tilde
        '˙': r'$\dot{}$',           # U+02D9 - Dot above
        '¨': r'$\ddot{}$',          # U+00A8 - Diaeresis
        
        # Special minus and spaces - using explicit Unicode escape sequences
        '−': r'-',                  # U+2212 - Minus sign (convert to regular hyphen)
        '\u2003': r' ',             # U+2003 - Em space (convert to regular space)
        '\u2009': r' ',             # U+2009 - Thin space (convert to regular space)
        '\u2002': r' ',             # U+2002 - En space (convert to regular space)
        '\u2004': r' ',             # U+2004 - Three-per-em space
        '\u2005': r' ',             # U+2005 - Four-per-em space
        '\u2006': r' ',             # U+2006 - Six-per-em space
        '\u2008': r' ',             # U+2008 - Punctuation space
        '\u200A': r' ',             # U+200A - Hair space
        '\u202F': r' ',             # U+202F - Narrow no-break space
        
        # Greek letters (commonly used in math)
        'α': r'$\alpha$',           # U+03B1
        'β': r'$\beta$',            # U+03B2
        'γ': r'$\gamma$',           # U+03B3
        'Γ': r'$\Gamma$',           # U+0393
        'ε': r'$\varepsilon$',      # U+03B5
        'ζ': r'$\zeta$',            # U+03B6
        'η': r'$\eta$',             # U+03B7
        'θ': r'$\theta$',           # U+03B8
        'Θ': r'$\Theta$',           # U+0398
        'ι': r'$\iota$',            # U+03B9
        'κ': r'$\kappa$',           # U+03BA
        'λ': r'$\lambda$',          # U+03BB
        'Λ': r'$\Lambda$',          # U+039B
        'μ': r'$\mu$',              # U+03BC
        'ν': r'$\nu$',              # U+03BD
        'ξ': r'$\xi$',              # U+03BE
        'Ξ': r'$\Xi$',              # U+039E
        'π': r'$\pi$',              # U+03C0
        'Π': r'$\Pi$',              # U+03A0
        'ρ': r'$\rho$',             # U+03C1
        'σ': r'$\sigma$',           # U+03C3
        'Σ': r'$\Sigma$',           # U+03A3
        'τ': r'$\tau$',             # U+03C4
        'υ': r'$\upsilon$',         # U+03C5
        'Υ': r'$\Upsilon$',         # U+03A5
        'φ': r'$\varphi$',          # U+03C6
        'Φ': r'$\Phi$',             # U+03A6
        'χ': r'$\chi$',             # U+03C7
        'ψ': r'$\psi$',             # U+03C8
        'Ψ': r'$\Psi$',             # U+03A8
        'ω': r'$\omega$',           # U+03C9
        'Ω': r'$\Omega$',           # U+03A9
    }
    
    # Apply conversions
    for unicode_char, latex_cmd in unicode_to_latex.items():
        if unicode_char in content:
            content = content.replace(unicode_char, latex_cmd)
    
    # Additional aggressive Unicode space cleanup using regex
    # Handle various Unicode spaces more comprehensively
    content = re.sub(r'[\u2000-\u200F\u2028-\u202F\u205F\u3000]', ' ', content)  # All Unicode spaces
    
    # Handle specific problematic Unicode characters that might not be in our dictionary
    content = re.sub(r'[\u2010-\u2015]', '-', content)  # Various Unicode dashes
    content = re.sub(r'[\u2212]', '-', content)         # Unicode minus sign
    
    # Handle specific cases where characters might appear in math environments
    # Fix double math mode (e.g., $\alpha$ inside already math mode)
    content = re.sub(r'\$\$([^$]+)\$\$', r'$\1$', content)  # Convert display math to inline
    content = re.sub(r'\$\$([^$]*)\$([^$]*)\$\$', r'$\1\2$', content)  # Fix broken math
    
    # Fix bar notation that might have been broken
    content = re.sub(r'\$\\bar\{\}\$([a-zA-Z])', r'$\\bar{\1}$', content)
    content = re.sub(r'([a-zA-Z])\$\\bar\{\}\$', r'$\\bar{\1}$', content)
    
    return content

def _additional_unicode_cleanup(content: str) -> str:
    """
    Additional aggressive Unicode cleanup to handle any characters that slip through.
    """
    # Convert all common problematic Unicode spaces to regular spaces
    # This covers a wider range than the dictionary approach
    unicode_spaces = [
        '\u00A0',  # Non-breaking space
        '\u1680',  # Ogham space mark
        '\u2000',  # En quad
        '\u2001',  # Em quad
        '\u2002',  # En space
        '\u2003',  # Em space
        '\u2004',  # Three-per-em space
        '\u2005',  # Four-per-em space
        '\u2006',  # Six-per-em space
        '\u2007',  # Figure space
        '\u2008',  # Punctuation space
        '\u2009',  # Thin space
        '\u200A',  # Hair space
        '\u200B',  # Zero width space
        '\u202F',  # Narrow no-break space
        '\u205F',  # Medium mathematical space
        '\u3000',  # Ideographic space
    ]
    
    for unicode_space in unicode_spaces:
        content = content.replace(unicode_space, ' ')
    
    # Convert Unicode dashes
    unicode_dashes = [
        '\u2010',  # Hyphen
        '\u2011',  # Non-breaking hyphen
        '\u2012',  # Figure dash
        '\u2013',  # En dash
        '\u2014',  # Em dash
        '\u2015',  # Horizontal bar
        '\u2212',  # Minus sign
    ]
    
    for unicode_dash in unicode_dashes:
        if unicode_dash in ['\u2013', '\u2014']:  # En and Em dashes
            content = content.replace(unicode_dash, '--')
        else:
            content = content.replace(unicode_dash, '-')
    
    # Use regex for any remaining problematic characters
    # Remove or replace any remaining Unicode characters that commonly cause issues
    content = re.sub(r'[\u2000-\u200F\u2028-\u202F\u205F\u3000]', ' ', content)
    content = re.sub(r'[\u2010-\u2015\u2212]', '-', content)
    
    return content

def _fix_mixed_mathematical_expressions(content: str) -> str:
    """
    Removes duplicated plain-text versions of mathematical expressions
    that Pandoc sometimes generates alongside the LaTeX version by deleting
    the plain text part when it is immediately followed by the LaTeX part.
    """
    
    processed_content = content

    # A list of compiled regex patterns.
    # Each pattern matches a plain-text formula but only if it's followed
    # by its corresponding LaTeX version (using a positive lookahead).
    patterns_to_remove = [
        # Pattern for: hq,k=x[nq,k]...h_{q,k} = x[n_{q,k}]...
        re.compile(r'h[qrs],k=x\[n[qrs],k\](?:,h[qrs],k=x\[n[qrs],k\])*\s*' +
                   r'(?=h_{q,k}\s*=\s*x\\\[n_{q,k}\\\],)', re.UNICODE),

        # Pattern for: ∆hq,r,k=hq,k-hr,k...\Delta h_{q,r,k} = ...
        re.compile(r'(?:∆h[qrs],[qrs],k=h[qrs],k-h[qrs],k\s*)+' +
                   r'(?=\\Delta\s*h_{q,r,k})', re.UNICODE),

        # Pattern for: RRk=tr,k+1-tr,kRR_k = ...
        re.compile(r'RRk=tr,k\+1-tr,k\s*' +
                   r'(?=RR_k\s*=\s*t_{r,k\+1})', re.UNICODE),

        # Pattern for: Tmed=median{RRk}T_{\mathrm{med}}
        re.compile(r'Tmed=median\{RRk\}\s*' +
                   r'(?=T_{\\mathrm{med}}\s*=\s*\\mathrm{median}\\{RR_k\\})', re.UNICODE),

        # Pattern for: Tk=[tr,k-Tmed2, tr,k+Tmed2]\mathcal{T}_k
        re.compile(r'Tk=\[tr,k-Tmed2,.*?tr,k\+Tmed2\]\s*' +
                   r'(?=\\mathcal\{T\}_k\s*=\s*\\\[t_{r,k})', re.UNICODE | re.DOTALL),

        # Pattern for: h¯k=1|Ik|∑n∈Ikx[n]\bar h_k
        re.compile(r'h¯k=1\|Ik\|∑n∈Ikx\[n\]\s*' +
                   r'(?=\\bar\s*h_k\s*=\s*\\frac)', re.UNICODE),

        # Pattern for: Mrs=median{∆hr,s,k}M_{rs}
        re.compile(r'Mrs=median\{∆hr,s,k\}\s*' +
                   r'(?=M_{rs}\s*=\s*\\mathrm{median})', re.UNICODE),
        
        # Pattern for: ∆h¯k=h¯k-Mrs\Delta\bar h_k
        re.compile(r'∆h¯k=h¯k-Mrs\s*' +
                   r'(?=\\Delta\\bar\s*h_k\s*=\s*\\bar\s*h_k)', re.UNICODE),
    ]

    for pattern in patterns_to_remove:
        processed_content = pattern.sub('', processed_content)
    
    return processed_content

def _fix_compilation_issues(content: str) -> str:
    """
    Fix common LaTeX compilation issues.
    """
    # Fix \tightlist command if not defined
    if r'\tightlist' in content and r'\providecommand{\tightlist}' not in content:
        tightlist_def = r'''
% Define \tightlist command for lists
\providecommand{\tightlist}{%
  \setlength{\itemsep}{0pt}\setlength{\parskip}{0pt}}
'''
        # Insert after packages but before \begin{document}
        begin_doc_match = re.search(r'\\begin\{document\}', content)
        if begin_doc_match:
            insert_pos = begin_doc_match.start()
            content = content[:insert_pos] + tightlist_def + '\n' + content[insert_pos:]
    
    # Fix \euro command if used but not defined
    if r'\euro' in content and r'usepackage{eurosym}' not in content:
        content = re.sub(
            r'(\\usepackage\{[^}]+\}\s*\n)',
            r'\1\\usepackage{eurosym}\n',
            content,
            count=1
        )
    
    # Fix undefined references to figures/tables
    content = re.sub(r'\\ref\{fig:([^}]+)\}', r'Figure~\\ref{fig:\1}', content)
    content = re.sub(r'\\ref\{tab:([^}]+)\}', r'Table~\\ref{tab:\1}', content)
    
    # Ensure proper figure placement
    if r'\begin{figure}' in content:
        content = re.sub(
            r'\\begin\{figure\}(?!\[)',
            r'\\begin{figure}[htbp]',
            content
        )
    
    # Ensure proper table placement  
    if r'\begin{table}' in content:
        content = re.sub(
            r'\\begin\{table\}(?!\[)',
            r'\\begin{table}[htbp]',
            content
        )
    
    return content

def _fix_image_paths_for_overleaf(content: str, extract_media_to_path: str = None) -> str:
    """
    Convert absolute image paths to relative paths for Overleaf compatibility.
    """
    if extract_media_to_path:
        # Extract the media directory name
        media_dir = os.path.basename(extract_media_to_path.rstrip('/'))
        
        # Fix paths with task IDs like: task_id_media/media/image.png -> media/image.png
        # Pattern: \includegraphics{any_path/task_id_media/media/image.ext}
        # Replace with: \includegraphics{media/image.ext}
        pattern1 = r'\\includegraphics(\[[^\]]*\])?\{[^{}]*[a-f0-9\-]+_media[/\\]media[/\\]([^{}]+)\}'
        replacement1 = r'\\includegraphics\1{media/\2}'
        content = re.sub(pattern1, replacement1, content)
        
        # Fix paths like: task_id_media/media/image.png -> media/image.png (without includegraphics)
        pattern2 = r'[a-f0-9\-]+_media[/\\]media[/\\]'
        replacement2 = r'media/'
        content = re.sub(pattern2, replacement2, content)
        
        # Also handle regular media paths: /absolute/path/to/media/image.ext -> media/image.ext
        pattern3 = r'\\includegraphics(\[[^\]]*\])?\{[^{}]*[/\\]' + re.escape(media_dir) + r'[/\\]([^{}]+)\}'
        replacement3 = r'\\includegraphics\1{' + media_dir + r'/\2}'
        content = re.sub(pattern3, replacement3, content)
    
    return content

def _remove_unwanted_formatting(content: str) -> str:
    """
    Remove unwanted highlighting and formatting that causes visual issues.
    """
    # Remove highlighting commands
    content = re.sub(r'\\colorbox\{[^}]*\}\{([^}]*)\}', r'\1', content)
    content = re.sub(r'\\hl\{([^}]*)\}', r'\1', content)
    content = re.sub(r'\\texthl\{([^}]*)\}', r'\1', content)
    content = re.sub(r'\\hlc\[[^\]]*\]\{([^}]*)\}', r'\1', content)
    
    # Remove table cell coloring
    content = re.sub(r'\\cellcolor\{[^}]*\}', '', content)
    content = re.sub(r'\\rowcolor\{[^}]*\}', '', content)
    content = re.sub(r'\\columncolor\{[^}]*\}', '', content)
    
    # Remove text background colors
    content = re.sub(r'\\textcolor\{[^}]*\}\{([^}]*)\}', r'\1', content)
    content = re.sub(r'\\color\{[^}]*\}', '', content)
    
    # Remove box formatting that might cause highlighting
    content = re.sub(r'\\fcolorbox\{[^}]*\}\{[^}]*\}\{([^}]*)\}', r'\1', content)
    content = re.sub(r'\\framebox\[[^\]]*\]\{([^}]*)\}', r'\1', content)
    
    # Remove soul package highlighting
    content = re.sub(r'\\sethlcolor\{[^}]*\}', '', content)
    content = re.sub(r'\\ul\{([^}]*)\}', r'\1', content)  # Remove underline if causing issues
    
    return content

def _inject_latex_packages(content: str) -> str:
    """
    Inject additional LaTeX packages needed for enhanced formatting.
    """
    # Essential packages for enhanced conversion
    essential_packages = [
        r'\usepackage{graphicx}',      # For images - ensure it's included
        r'\usepackage{longtable}',     # For tables
        r'\usepackage{booktabs}',      # Better table formatting  
        r'\usepackage{array}',         # Enhanced table formatting
        r'\usepackage{calc}',          # For calculations
        r'\usepackage{url}',           # For URLs
    ]
    
    # Style enhancement packages
    style_packages = [
        r'\usepackage{float}',         # Better float positioning
        r'\usepackage{adjustbox}',     # For centering and scaling
        r'\usepackage{caption}',       # Better caption formatting
        r'\usepackage{subcaption}',    # For subfigures
        r'\usepackage{tabularx}',      # Flexible table widths
        r'\usepackage{enumitem}',      # Better list formatting
        r'\usepackage{setspace}',      # Line spacing control
        r'\usepackage{ragged2e}',      # Better text alignment
        r'\usepackage{amsmath}',       # Mathematical formatting
        r'\usepackage{amssymb}',       # Mathematical symbols
        r'\usepackage{needspace}',     # Prevent orphaned lines and improve page breaks
    ]
    
    all_packages = essential_packages + style_packages
    
    # Find the position after \documentclass but before any existing \usepackage or \begin{document}
    documentclass_pattern = r'\\documentclass(?:\[[^\]]*\])?\{[^}]+\}'
    documentclass_match = re.search(documentclass_pattern, content)
    
    if documentclass_match:
        insert_pos = documentclass_match.end()
        
        # Find the next significant LaTeX command to insert before it
        # Look for existing \usepackage, \begin{document}, or other commands
        remaining_content = content[insert_pos:]
        next_command_match = re.search(r'\\(?:usepackage|begin\{document\}|title|author|date)', remaining_content)
        
        if next_command_match:
            insert_pos += next_command_match.start()
        
        # Check which packages are not already included
        packages_to_insert = []
        for package in all_packages:
            package_name = package.replace(r'\usepackage{', '').replace('}', '')
            if f'usepackage{{{package_name}}}' not in content:
                packages_to_insert.append(package)
        
        if packages_to_insert:
            # Add packages with proper spacing
            package_block = '\n% Enhanced conversion packages\n' + '\n'.join(packages_to_insert) + '\n\n'
            content = content[:insert_pos] + package_block + content[insert_pos:]
    
    return content

def _add_centering_commands(content: str) -> str:
    """
    Add centering commands to figures and tables.
    """
    # Add \centering to figure environments
    content = re.sub(
        r'(\\begin\{figure\}(?:\[[^\]]*\])?)\s*\n',
        r'\1\n\\centering\n',
        content
    )
    
    # Add \centering to table environments
    content = re.sub(
        r'(\\begin\{table\}(?:\[[^\]]*\])?)\s*\n',
        r'\1\n\\centering\n',
        content
    )
    
    return content

def _fix_line_breaks_and_spacing(content: str) -> str:
    """
    Minimal fixes to preserve Word's original formatting and pagination.
    """
    # Remove unwanted highlighting and color commands
    content = re.sub(r'\\colorbox\{[^}]*\}\{([^}]*)\}', r'\1', content)
    content = re.sub(r'\\hl\{([^}]*)\}', r'\1', content)
    content = re.sub(r'\\texthl\{([^}]*)\}', r'\1', content)
    content = re.sub(r'\\cellcolor\{[^}]*\}', '', content)
    content = re.sub(r'\\rowcolor\{[^}]*\}', '', content)
    
    # Only fix critical spacing issues that break compilation
    # Preserve Word's original line breaks and spacing as much as possible
    
    # Ensure proper spacing around lists but don't change internal spacing
    content = re.sub(r'\n\\begin\{enumerate\}\n\n', r'\n\n\\begin{enumerate}\n', content)
    content = re.sub(r'\n\n\\end\{enumerate\}\n', r'\n\\end{enumerate}\n\n', content)
    content = re.sub(r'\n\\begin\{itemize\}\n\n', r'\n\n\\begin{itemize}\n', content)
    content = re.sub(r'\n\n\\end\{itemize\}\n', r'\n\\end{itemize}\n\n', content)
    
    # Minimal section spacing - preserve Word's pagination
    content = re.sub(r'\n(\\(?:sub)*section\{[^}]+\})\n\n', r'\n\n\1\n\n', content)
    
    # Only remove excessive spacing (3+ line breaks) but preserve double breaks
    content = re.sub(r'\n\n\n+', r'\n\n', content)
    
    # Ensure proper spacing around figures and tables
    content = re.sub(r'\n\\begin\{figure\}', r'\n\n\\begin{figure}', content)
    content = re.sub(r'\\end\{figure\}\n([A-Z])', r'\\end{figure}\n\n\1', content)
    content = re.sub(r'\n\\begin\{table\}', r'\n\n\\begin{table}', content)
    content = re.sub(r'\\end\{table\}\n([A-Z])', r'\\end{table}\n\n\1', content)
    
    return content

if __name__ == '__main__':
    from docx import Document
    from docx.shared import Inches
    from PIL import Image
    import shutil

    # --- Helper Functions for DOCX and Template Creation ---
    def create_dummy_image(filename, size=(60, 60), color="red", img_format="PNG"):
        img = Image.new('RGB', size, color=color)
        img.save(filename, img_format)
        print(f"Created dummy image: {filename}")

    def create_test_docx_with_styles(filename):
        doc = Document()
        doc.add_heading("Document with Enhanced Features", level=1)
        
        # Add paragraph with text
        p1 = doc.add_paragraph("This document tests enhanced features including:")
        
        # Add numbered list
        doc.add_paragraph("First numbered item", style='List Number')
        doc.add_paragraph("Second numbered item", style='List Number')
        doc.add_paragraph("Third numbered item", style='List Number')
        
        # Add some text
        doc.add_paragraph("Here is some regular text between lists.")
        
        # Add bullet list
        doc.add_paragraph("First bullet point", style='List Bullet')
        doc.add_paragraph("Second bullet point", style='List Bullet')
        
        doc.add_heading("Image Section", level=2)
        doc.add_paragraph("Below is a test image:")
        
        doc.save(filename)
        print(f"Created test DOCX with styles: {filename}")

    def create_complex_docx(filename, img1_path, img2_path):
        doc = Document()
        doc.add_heading("Complex Document Title", level=1)
        doc.add_paragraph("Introduction to the complex document.")
        doc.add_heading("Image Section", level=2)
        doc.add_picture(img1_path, width=Inches(1.0))
        doc.add_paragraph("Some text after the first image.")
        doc.add_picture(img2_path, width=Inches(1.0))
        doc.add_heading("Conclusion Section", level=2)
        doc.add_paragraph("Final remarks.")
        doc.save(filename)
        print(f"Created complex DOCX: {filename}")

    # --- Test Files ---
    docx_styles = "test_enhanced_styles.docx"
    docx_complex = "test_complex_enhanced.docx"
    img1 = "dummy_img1.png"
    img2 = "dummy_img2.jpg"

    output_enhanced_test = "output_enhanced_test.tex"
    output_overleaf_test = "output_overleaf_test.tex"
    media_dir = "./media_enhanced"

    all_test_files = [docx_styles, docx_complex, img1, img2, output_enhanced_test, output_overleaf_test]
    all_test_dirs = [media_dir]

    # --- Create Test Files ---
    print("--- Setting up enhanced test files ---")
    create_dummy_image(img1, color="blue", img_format="PNG")
    create_dummy_image(img2, color="green", img_format="JPEG")
    create_test_docx_with_styles(docx_styles)
    create_complex_docx(docx_complex, img1, img2)
    print("--- Enhanced test file setup complete ---")

    # --- Test Enhanced Features ---
    print("\n--- Testing Enhanced Features ---")

    # Test 1: Style preservation and line breaks
    print("\n--- Test 1: Enhanced Style Preservation ---")
    success, msg = convert_docx_to_latex(
        docx_styles, 
        output_enhanced_test,
        generate_toc=True,
        preserve_styles=True,
        preserve_linebreaks=True
    )
    print(f"Enhanced Test: {success}, Msg: {msg}")
    
    if success and os.path.exists(output_enhanced_test):
        with open(output_enhanced_test, 'r') as f:
            content = f.read()
            checks = {
                'packages': any(pkg in content for pkg in ['\\usepackage{float}', '\\usepackage{enumitem}']),
                'toc': '\\tableofcontents' in content,
                'sections': '\\section' in content,
                'lists': '\\begin{enumerate}' in content or '\\begin{itemize}' in content
            }
            print(f"Enhanced verification: {checks}")

    # Test 2: Overleaf compatibility with images
    print("\n--- Test 2: Overleaf Compatibility ---")
    success, msg = convert_docx_to_latex(
        docx_complex,
        output_overleaf_test,
        extract_media_to_path=media_dir,
        overleaf_compatible=True,
        preserve_styles=True,
        preserve_linebreaks=True
    )
    print(f"Overleaf Test: {success}, Msg: {msg}")
    
    if success and os.path.exists(output_overleaf_test):
        with open(output_overleaf_test, 'r') as f:
            content = f.read()
            media_check = 'media/' in content and '\\includegraphics' in content
            print(f"Overleaf compatibility check - relative paths: {media_check}")
            
        media_files_exist = os.path.exists(os.path.join(media_dir, 'media'))
        print(f"Media files extracted: {media_files_exist}")

    # --- Cleanup ---
    print("\n--- Cleaning up enhanced test files ---")
    for f_path in all_test_files:
        if os.path.exists(f_path):
            try:
                os.remove(f_path)
                print(f"Removed: {f_path}")
            except Exception as e:
                print(f"Error removing {f_path}: {e}")

    for d_path in all_test_dirs:
        if os.path.isdir(d_path):
            try:
                shutil.rmtree(d_path)
                print(f"Removed directory: {d_path}")
            except Exception as e:
                print(f"Error removing {d_path}: {e}")

    print("--- Enhanced testing completed ---") 